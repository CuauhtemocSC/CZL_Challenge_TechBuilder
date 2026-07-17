import os
import io
from docx import Document
from PIL import Image
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document as LCDocument
from transformers import BlipProcessor, BlipForConditionalGeneration

def inicializar_modelo_visual():
    print("📸 Cargando modelo de visión artificial (BLIP) en CPU...")
    # Modelo gratuito de subtitulado de imágenes de Salesforce
    processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
    model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
    return processor, model

def describir_imagen(imagen_bytes, processor, model):
    try:
        imagen = Image.open(io.BytesIO(imagen_bytes)).convert('RGB')
        # Pedimos al modelo que describa la imagen de forma condicional
        inputs = processor(imagen, text="A screenshot or diagram showing", return_tensors="pt")
        out = model.generate(**inputs)
        descripcion = processor.decode(out[0], skip_special_tokens=True)
        return f"[IMAGEN ILUSTRATIVA: {descripcion}]"
    except Exception as e:
        print(f"⚠️ No se pudo procesar una imagen: {e}")
        return ""

def procesar_documentos_multimodal():
    processor, model = inicializar_modelo_visual()
    
    carpeta_docs = "./documentos"
    documentos_para_vectorizar = []
    
    if not os.path.exists(carpeta_docs) or not os.listdir(carpeta_docs):
        print("⚠️ Asegúrate de tener archivos .docx en la carpeta './documentos'")
        return

    print("📂 Leyendo documentos e indexando imágenes...")
    for archivo in os.listdir(carpeta_docs):
        if archivo.endswith(".docx"):
            ruta_completa = os.path.join(carpeta_docs, archivo)
            print(f"📖 Analizando: {archivo}")
            
            # Abrimos el documento a bajo nivel con python-docx
            doc = Document(ruta_completa)
            texto_acumulado = []
            
            # Recorremos los párrafos del documento
            for para in doc.paragraphs:
                texto_acumulado.append(para.text)
                
                # Buscamos si el párrafo contiene elementos gráficos/imágenes
                if 'graphic' in para._p.xml:
                    for rel_id, rel in doc.part.rels.items():
                        if "image" in rel.target_ref:
                            try:
                                # Extraemos los bytes de la imagen incrustada
                                img_bytes = rel.target_part.blob
                                descripcion = describir_imagen(img_bytes, processor, model)
                                if descripcion:
                                    print(f"   👁️ Imagen interpretada: {descripcion}")
                                    texto_acumulado.append(f"\n{descripcion}\n")
                            except:
                                pass
                                
            texto_final = "\n".join(texto_acumulado)
            # Creamos un documento compatible con LangChain
            documentos_para_vectorizar.append(
                LCDocument(page_content=texto_final, metadata={"source": archivo})
            )

    # Reutilizamos nuestro particionamiento y base de datos
    print("✂️ Dividiendo texto y descripciones en fragmentos...")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    fragmentos = text_splitter.split_documents(documentos_para_vectorizar)
    
    print("🤖 Inicializando modelo de embeddings...")
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    
    print("💾 Guardando base de datos multimodal en './vector_db'...")
    Chroma.from_documents(
        documents=fragmentos,
        embedding=embeddings,
        persist_directory="./vector_db"
    )
    print("✅ ¡Base de datos vectorial multimodal actualizada con éxito!")

if __name__ == "__main__":
    procesar_documentos_multimodal()