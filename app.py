import os
import io
import streamlit as st
from docx import Document
from PIL import Image
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document as LCDocument
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_community.llms import HuggingFacePipeline
from transformers import pipeline, AutoModelForCausalLM, AutoTokenizer, BlipProcessor, BlipForConditionalGeneration

st.set_page_config(page_title="Asistente Operativo Multimodal", page_icon="🤖", layout="centered")

# --- PROCESAMIENTO AUTOMÁTICO EN LA NUBE ---
@st.cache_resource
def procesar_e_inicializar_todo():
    carpeta_docs = "./documentos"
    ruta_db = "./vector_db_cloud"
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    
    # Si la base de datos no existe en la nube, procesamos los Word en el arranque
    if not os.path.exists(ruta_db):
        st.info("📦 Inicializando el cerebro del agente por primera vez en la nube...")
        
        # Cargar modelo visual para las imágenes
        processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
        model_v = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
        
        documentos_para_vectorizar = []
        if os.path.exists(carpeta_docs):
            for archivo in os.listdir(carpeta_docs):
                if archivo.endswith(".docx"):
                    doc = Document(os.path.join(carpeta_docs, archivo))
                    texto_acumulado = []
                    for para in doc.paragraphs:
                        texto_acumulado.append(para.text)
                        if 'graphic' in para._p.xml:
                            for rel_id, rel in doc.part.rels.items():
                                if "image" in rel.target_ref:
                                    try:
                                        img_bytes = rel.target_part.blob
                                        imagen = Image.open(io.BytesIO(img_bytes)).convert('RGB')
                                        inputs = processor(imagen, text="A diagram showing", return_tensors="pt")
                                        out = model_v.generate(**inputs)
                                        desc = processor.decode(out[0], skip_special_tokens=True)
                                        texto_acumulado.append(f"\n[IMAGEN: {desc}]\n")
                                    except: pass
                    
                    documentos_para_vectorizar.append(
                        LCDocument(page_content="\n".join(texto_acumulado), metadata={"source": archivo})
                    )
        
        if documentos_para_vectorizar:
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
            fragmentos = text_splitter.split_documents(documentos_para_vectorizar)
            vector_store = Chroma.from_documents(documents=fragmentos, embedding=embeddings, persist_directory=ruta_db)
        else:
            # Si no hay documentos, creamos una BD vacía para evitar errores
            vector_store = Chroma.from_documents(documents=[LCDocument(page_content="No hay datos.")], embedding=embeddings, persist_directory=ruta_db)
    else:
        vector_store = Chroma(persist_directory=ruta_db, embedding_function=embeddings)
        
    retriever = vector_store.as_retriever(search_kwargs={"k": 3})

    # Cargar LLM
    model_id = "Qwen/Qwen2.5-0.5B-Instruct"
    tokenizer = AutoTokenizer.from_pretrained(model_id)
    model_llm = AutoModelForCausalLM.from_pretrained(model_id)
    
    pipe = pipeline("text-generation", model=model_llm, tokenizer=tokenizer, max_new_tokens=256, temperature=0.2)
    llm = HuggingFacePipeline(pipeline=pipe)

    system_prompt = (
        "Eres un asistente operativo experto. Responde en español basándote ÚNICAMENTE en el contexto provisto. "
        "Si la información no está en el contexto, di amablemente que no la encuentras en los manuales.\n\n"
        "Contexto:\n{context}\n\nPregunta: {input}\nRespuesta:"
    )
    prompt = ChatPromptTemplate.from_template(system_prompt)
    
    def format_docs(docs):
        return "\n\n".join(doc.page_content for doc in docs)

    return (
        {"context": retriever | format_docs, "input": RunnablePassthrough()}
        | prompt | llm | StrOutputParser()
    )

# --- INTERFAZ ---
st.title("🤖 Asistente Operativo Corporativo")
st.markdown("---")

agente = procesar_e_inicializar_todo()

if "messages" not in st.session_state:
    st.session_state.messages = [{"role": "assistant", "content": "¡Hola! He procesado las guías operativas. ¿Qué duda deseas resolver?"}]

for message in st.session_state.messages:
    with st.chat_message(message["role"]): st.write(message["content"])

if pregunta_usuario := st.chat_input("Escribe tu duda aquí..."):
    with st.chat_message("user"): st.write(pregunta_usuario)
    st.session_state.messages.append({"role": "user", "content": pregunta_usuario})

    with st.chat_message("assistant"):
        with st.spinner("Analizando manuales..."):
            respuesta = agente.invoke(pregunta_usuario)
            st.write(respuesta)
    st.session_state.messages.append({"role": "assistant", "content": respuesta})